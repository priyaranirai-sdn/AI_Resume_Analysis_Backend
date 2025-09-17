import os
import re
import json
from typing import List, Dict, Any, Tuple
import PyPDF2
import docx
from docx import Document
import io

class ResumeAnalyzer:
    def __init__(self):
        # Initialize AI models lazily to avoid startup delays
        self.similarity_pipeline = None
        self.text_classification_pipeline = None
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF, DOC, or DOCX files"""
        try:
            if not file_content:
                raise ValueError("File content is empty")
            
            file_extension = filename.lower().split('.')[-1]
            print(f"Processing file: {filename}, extension: {file_extension}, size: {len(file_content)} bytes")
            
            if file_extension == 'pdf':
                text = self._extract_from_pdf(file_content)
            elif file_extension in ['doc', 'docx']:
                text = self._extract_from_docx(file_content)
            elif file_extension == 'txt':
                text = self._extract_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: PDF, DOC, DOCX, TXT")
            
            if not text or not text.strip():
                raise ValueError("No text content extracted from file")
            
            print(f"Successfully extracted {len(text)} characters from {filename}")
            return text.strip()
            
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            raise Exception(f"Error extracting text from {filename}: {str(e)}")
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            
            if len(pdf_reader.pages) == 0:
                raise Exception("PDF has no pages")
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        print(f"Warning: No text extracted from page {i+1}")
                except Exception as page_error:
                    print(f"Warning: Error extracting text from page {i+1}: {page_error}")
                    continue
            
            if not text.strip():
                raise Exception("No text content found in PDF")
                
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            
            print(f"DOCX file has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            # Extract text from paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"
                    print(f"Paragraph {i+1}: {para_text[:50]}...")
                else:
                    print(f"Paragraph {i+1}: Empty")
            
            # Also extract text from tables
            for table_idx, table in enumerate(doc.tables):
                print(f"Table {table_idx+1} has {len(table.rows)} rows")
                for row_idx, row in enumerate(table.rows):
                    row_text = ""
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text += cell_text + " "
                    if row_text.strip():
                        text += row_text.strip() + "\n"
                        print(f"Table {table_idx+1}, Row {row_idx+1}: {row_text.strip()[:50]}...")
            
            # Try to extract text from headers and footers
            try:
                for section in doc.sections:
                    if section.header:
                        header_text = section.header.paragraphs[0].text.strip()
                        if header_text:
                            text += header_text + "\n"
                            print(f"Header: {header_text}")
                    
                    if section.footer:
                        footer_text = section.footer.paragraphs[0].text.strip()
                        if footer_text:
                            text += footer_text + "\n"
                            print(f"Footer: {footer_text}")
            except Exception as header_error:
                print(f"Warning: Could not extract headers/footers: {header_error}")
            
            print(f"Total extracted text length: {len(text)} characters")
            
            if not text.strip():
                # Try alternative extraction method
                print("Trying alternative DOCX extraction...")
                text = self._extract_from_docx_alternative(file_content)
                if not text.strip():
                    raise Exception("No text content found in DOCX file using any method")
                
            return text.strip()
        except Exception as e:
            print(f"Error in primary DOCX extraction: {e}")
            # Try alternative method
            try:
                text = self._extract_from_docx_alternative(file_content)
                if text.strip():
                    return text.strip()
            except Exception as alt_error:
                print(f"Alternative extraction also failed: {alt_error}")
            
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_from_docx_alternative(self, file_content: bytes) -> str:
        """Alternative method to extract text from DOCX using zipfile"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            text = ""
            
            # DOCX files are ZIP archives
            with zipfile.ZipFile(io.BytesIO(file_content)) as docx_zip:
                # Read the main document XML
                if 'word/document.xml' in docx_zip.namelist():
                    doc_xml = docx_zip.read('word/document.xml')
                    root = ET.fromstring(doc_xml)
                    
                    # Extract text from all text nodes
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text += elem.text.strip() + " "
                        if elem.tail and elem.tail.strip():
                            text += elem.tail.strip() + " "
            
            return text.strip()
        except Exception as e:
            print(f"Alternative DOCX extraction failed: {e}")
            return ""
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    if text.strip():
                        return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            if not text.strip():
                raise Exception("No text content found in TXT file")
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def extract_skills_from_text(self, text: str) -> List[str]:
        """Extract skills from resume text using pattern matching"""
        # Common technical skills patterns
        skill_patterns = [
            r'\b(?:JavaScript|JS|React|Angular|Vue|Node\.?js|Express|jQuery)\b',
            r'\b(?:Python|Java|C\+\+|C#|PHP|Ruby|Go|Rust|Swift|Kotlin)\b',
            r'\b(?:HTML|CSS|SASS|SCSS|Bootstrap|Tailwind)\b',
            r'\b(?:SQL|MySQL|PostgreSQL|MongoDB|Redis|Elasticsearch)\b',
            r'\b(?:AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Git|GitHub|GitLab)\b',
            r'\b(?:MERN|MEAN|LAMP|Django|Flask|Spring|Laravel|Rails)\b',
            r'\b(?:Machine Learning|ML|AI|Data Science|Analytics|Big Data)\b',
            r'\b(?:Agile|Scrum|DevOps|CI/CD|Microservices|REST|API)\b',
            r'\b(?:Linux|Unix|Windows|macOS|iOS|Android)\b',
            r'\b(?:Photoshop|Illustrator|Figma|Sketch|Adobe|Design)\b'
        ]
        
        skills = set()
        text_lower = text.lower()
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.update([match.strip() for match in matches])
        
        # Also look for skills mentioned in common formats
        skill_sections = re.findall(r'(?:skills?|technologies?|tools?|languages?)[:\-]?\s*([^.\n]+)', text_lower)
        for section in skill_sections:
            # Split by common separators
            section_skills = re.split(r'[,;|•\-\n]', section)
            for skill in section_skills:
                skill = skill.strip()
                if len(skill) > 2 and len(skill) < 50:  # Reasonable skill length
                    skills.add(skill.title())
        
        return list(skills)
    
    def extract_experience_years(self, text: str) -> int:
        """Extract years of experience from resume text"""
        # Look for experience patterns
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'(?:experience|exp)[:\-]?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*(?:in|of)',
            r'(\d+)\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:professional|work)',
            r'(\d+)\+?\s*years?\s*(?:in\s*)?(?:software|development|programming)',
            r'(\d+)\+?\s*years?\s*(?:as\s*)?(?:developer|engineer|programmer)',
        ]
        
        max_years = 0
        text_lower = text.lower()
        
        print(f"Looking for experience in text: {text_lower[:200]}...")
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text_lower)
            print(f"Pattern '{pattern}' found matches: {matches}")
            for match in matches:
                try:
                    years = int(match)
                    max_years = max(max_years, years)
                    print(f"Found {years} years of experience")
                except ValueError:
                    continue
        
        # Also look for simple patterns like "4 years" or "3+ years"
        simple_patterns = [
            r'(\d+)\+?\s*years?',
            r'(\d+)\s*years?',
        ]
        
        for pattern in simple_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match)
                    # Only consider reasonable experience years (1-50)
                    if 1 <= years <= 50:
                        max_years = max(max_years, years)
                        print(f"Found {years} years in simple pattern")
                except ValueError:
                    continue
        
        print(f"Final experience extracted: {max_years} years")
        return max_years
    
    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity between two texts using sentence transformers"""
        try:
            # Lazy load the similarity pipeline
            if self.similarity_pipeline is None:
                from transformers import pipeline
                self.similarity_pipeline = pipeline(
                    "feature-extraction", 
                    model="sentence-transformers/all-MiniLM-L6-v2"
                )
            
            # Get embeddings for both texts
            embeddings1 = self.similarity_pipeline(text1)
            embeddings2 = self.similarity_pipeline(text2)
            
            # Calculate cosine similarity
            import numpy as np
            from sklearn.metrics.pairwise import cosine_similarity
            
            # Convert to numpy arrays and calculate similarity
            emb1 = np.array(embeddings1).mean(axis=0).reshape(1, -1)
            emb2 = np.array(embeddings2).mean(axis=0).reshape(1, -1)
            
            similarity = cosine_similarity(emb1, emb2)[0][0]
            return float(similarity)
        except Exception as e:
            print(f"Error calculating similarity: {e}")
            # Fallback to simple text similarity
            return self._simple_text_similarity(text1, text2)
    
    def _simple_text_similarity(self, text1: str, text2: str) -> float:
        """Simple text similarity based on common words"""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def analyze_resume_match(self, resume_text: str, job_description: str, required_skills: List[str], 
                           required_experience: int) -> Dict[str, Any]:
        """Analyze how well a resume matches a job description"""
        
        # Extract information from resume
        resume_skills = self.extract_skills_from_text(resume_text)
        resume_experience = self.extract_experience_years(resume_text)
        
        # Calculate skill matches
        required_skills_lower = [skill.lower() for skill in required_skills]
        resume_skills_lower = [skill.lower() for skill in resume_skills]
        
        matched_skills = []
        missing_skills = []
        
        for skill in required_skills_lower:
            # Check for exact match or partial match
            found = False
            for resume_skill in resume_skills_lower:
                if skill in resume_skill or resume_skill in skill:
                    matched_skills.append(skill)
                    found = True
                    break
            if not found:
                missing_skills.append(skill)
        
        # Calculate skill match percentage
        skill_match_percentage = (len(matched_skills) / len(required_skills)) * 100 if required_skills else 0
        
        # Check experience match
        experience_match = resume_experience >= required_experience
        
        # Calculate overall similarity using AI
        similarity_score = self.calculate_similarity(resume_text, job_description)
        
        # Calculate overall match percentage
        # Weight: 40% skills, 30% experience, 30% content similarity
        overall_match = (
            skill_match_percentage * 0.4 +
            (100 if experience_match else 0) * 0.3 +
            similarity_score * 100 * 0.3
        )
        
        # Determine match status
        if overall_match >= 80:
            match_status = "Match"
            suitability_rating = "High"
        elif overall_match >= 60:
            match_status = "Partial Match"
            suitability_rating = "Medium"
        else:
            match_status = "Not a Match"
            suitability_rating = "Low"
        
        # Generate gaps analysis
        gaps_analysis = self._generate_gaps_analysis(matched_skills, missing_skills, 
                                                   experience_match, required_experience, resume_experience)
        
        return {
            "match_percentage": round(overall_match, 2),
            "confidence_score": round(similarity_score * 100, 2),
            "is_match": match_status,
            "suitability_rating": suitability_rating,
            "skills_match": {
                "matched_skills": matched_skills,
                "missing_skills": missing_skills,
                "match_percentage": round(skill_match_percentage, 2)
            },
            "experience_match": experience_match,
            "resume_experience": resume_experience,
            "required_experience": required_experience,
            "gaps_analysis": gaps_analysis,
            "resume_skills": resume_skills
        }
    
    def _generate_gaps_analysis(self, matched_skills: List[str], missing_skills: List[str], 
                              experience_match: bool, required_exp: int, resume_exp: int) -> str:
        """Generate detailed gaps analysis"""
        gaps = []
        
        if missing_skills:
            gaps.append(f"Missing required skills: {', '.join(missing_skills)}")
        
        if not experience_match:
            gaps.append(f"Experience gap: Resume shows {resume_exp} years, but {required_exp}+ years required")
        
        if not gaps:
            return "No significant gaps identified. Candidate meets all requirements."
        
        return "Gaps identified: " + "; ".join(gaps)

# Global instance
resume_analyzer = ResumeAnalyzer()
